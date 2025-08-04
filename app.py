import streamlit as st
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import os
import time
import re
from pathlib import Path
from dotenv import load_dotenv
from typing import Optional, Tuple, Dict, Any, List
import logging
import base64
import io

try:
    from ethics_handler import render_ethics_chat_interface, initialize_ethics_session_state
    ETHICS_AVAILABLE = True
except ImportError as e:
    ETHICS_AVAILABLE = False

# Import our localization system
from localization import language_manager, t, init_language_system, render_language_selector, get_rtl_css, get_language_specific_ai_prompt

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configuration - Updated for guided flow
class Config:
    PROJECT_NAME = "Roehampton University Chatbot"
    COMPANY_NAME = "University of Roehampton"
    DATA_FOLDER = "data"
    AUDIO_FOLDER = "audio_responses"
    TEMP_AUDIO_FOLDER = "temp_audio"
    LOGO_PATH = "logo.png"
    STUDENT_DATA_FILE = "student_modules_with_pdfs.xlsx"  # Excel file with student data
    ETHICS_FOLDER = "ethics_documents"  # Folder for ethics documents
    MAX_TOKENS = 1500
    TEMPERATURE = 0.3
    MODEL = "gpt-3.5-turbo"
    TTS_MODEL = "tts-1"
    TTS_VOICE = "alloy"
    MAX_CONTENT_LENGTH = 15000
    PREVIEW_LENGTH = 800
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx']
    SUPPORTED_VOICES = {
        'alloy': 'Alloy (Neutral)',
        'echo': 'Echo (Male)', 
        'fable': 'Fable (British Male)',
        'onyx': 'Onyx (Deep Male)',
        'nova': 'Nova (Female)',
        'shimmer': 'Shimmer (Soft Female)'
    }

# Initialize OpenAI
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if OPENAI_API_KEY:
    client = OpenAI(api_key=OPENAI_API_KEY)
else:
    client = None


def load_logo_from_assets() -> Optional[str]:
    """Load logo from assets folder and encode as base64"""
    # Try multiple possible logo locations in assets folder
    possible_paths = [
        Path("assets") / "logo.png",
        Path("assets") / "logo.jpg", 
        Path("assets") / "logo.jpeg",
        Path("assets") / "logo.svg",
        Path("assets") / "roehampton_logo.png",
        Path("assets") / "university_logo.png"
    ]
    
    for logo_path in possible_paths:
        if logo_path.exists():
            try:
                with open(logo_path, "rb") as img_file:
                    img_bytes = img_file.read()
                    img_base64 = base64.b64encode(img_bytes).decode()
                    logger.info(f"Successfully loaded logo from: {logo_path}")
                    return img_base64
            except Exception as e:
                logger.warning(f"Error loading logo from {logo_path}: {e}")
                continue
    
    logger.info("No logo found in assets folder")
    return None 

# Page configuration
st.set_page_config(
    page_title=Config.PROJECT_NAME,
    page_icon= "🎓",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': f"# {Config.PROJECT_NAME}\nGuided coursework assistant for University of Roehampton students"
    }
)

def initialize_session_state() -> None:
    """Initialize session state for guided conversation flow"""
    # Conversation flow states
    if 'conversation_step' not in st.session_state:
        st.session_state.conversation_step = 'welcome'  # welcome -> path_selection -> student_id -> code -> module -> coursework -> chat
    
    # Authentication data
    if 'student_id' not in st.session_state:
        st.session_state.student_id = None
    if 'student_code' not in st.session_state:
        st.session_state.student_code = None
    if 'student_data' not in st.session_state:
        st.session_state.student_data = None
    if 'selected_path' not in st.session_state:
        st.session_state.selected_path = None  # 'ethics' or 'coursework'
    
    # Module and coursework selection
    if 'available_modules' not in st.session_state:
        st.session_state.available_modules = []
    if 'selected_module' not in st.session_state:
        st.session_state.selected_module = None
    if 'selected_coursework' not in st.session_state:
        st.session_state.selected_coursework = None
    if 'current_document' not in st.session_state:
        st.session_state.current_document = None
    
    # Chat data
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'audio_enabled' not in st.session_state:
        st.session_state.audio_enabled = True
    if 'selected_voice' not in st.session_state:
        st.session_state.selected_voice = Config.TTS_VOICE
    if 'audio_responses' not in st.session_state:
        st.session_state.audio_responses = {}
    
    # Error handling
    if 'error_message' not in st.session_state:
        st.session_state.error_message = None
    if 'retry_count' not in st.session_state:
        st.session_state.retry_count = 0
    
    # Student database
    if 'student_database' not in st.session_state:
        st.session_state.student_database = None
    if 'database_loaded' not in st.session_state:
        st.session_state.database_loaded = False

      # Add ethics-specific initialization
    if 'selected_ethics_category' not in st.session_state:
        st.session_state.selected_ethics_category = None
    if 'ethics_document' not in st.session_state:
        st.session_state.ethics_document = None
    
    # Initialize ethics session state if available
    if ETHICS_AVAILABLE:
        initialize_ethics_session_state()
    
    # Initialize language system
    init_language_system()

def extract_coursework_type_from_filename(pdf_filename: str) -> str:
    """Extract coursework type from PDF filename"""
    filename_lower = pdf_filename.lower()
    
    # Machine Learning specific patterns
    if 'coursework1' in filename_lower:
        return 'Coursework 1'
    elif 'coursework2' in filename_lower:
        return 'Coursework 2'
    elif 'coursework3' in filename_lower:
        return 'Coursework 3'
    elif 'assignment1' in filename_lower:
        return 'Assignment 1'
    elif 'assignment2' in filename_lower:
        return 'Assignment 2'
    elif 'assignment3' in filename_lower:
        return 'Assignment 3'
    else:
        return 'Course Materials'
    
def format_pdf_display_name(pdf_filename: str) -> str:
    """Create user-friendly display name from PDF filename"""
    # Remove extension and replace underscores
    name = Path(pdf_filename).stem.replace('_', ' ')
    
    # Capitalize each word
    name = ' '.join(word.capitalize() for word in name.split())
    
    return name

@st.cache_data(show_spinner=False)
def load_student_database() -> Tuple[Optional[Dict], str]:
    """Load student database from Excel file with support for multiple PDFs per module"""
    try:
        excel_path = Path(Config.STUDENT_DATA_FILE)
        if not excel_path.exists():
            return None, f"Student database file not found: {Config.STUDENT_DATA_FILE}"
        
        # Read Excel file
        df = pd.read_excel(excel_path)
        
        # Validate required columns
        required_columns = ['Student ID', 'Code', 'Programme', 'Module', 'PDF File']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            return None, f"Missing required columns in Excel file: {missing_columns}"
        
        # Create enhanced data structures to handle multiple PDFs per module
        student_database = {
            'students': {},  # student_id -> {code, programme, modules}
            'student_codes': {},  # student_id -> code
            'student_modules': {},  # student_id -> {module_name: [pdf_files]}
            'programme_modules': {}  # programme -> {module_name: [pdf_files]}
        }
        
        for _, row in df.iterrows():
            student_id = str(row['Student ID'])
            code = int(row['Code'])
            programme = row['Programme']
            module = row['Module']
            pdf_file = row['PDF File']
            
            # Initialize student if not exists
            if student_id not in student_database['students']:
                student_database['students'][student_id] = {
                    'code': code,
                    'programme': programme,
                    'modules': {}
                }
                student_database['student_codes'][student_id] = code
                student_database['student_modules'][student_id] = {}
            
            # Initialize module if not exists for this student
            if module not in student_database['student_modules'][student_id]:
                student_database['student_modules'][student_id][module] = []
            
            # Create PDF data with enhanced info for Machine Learning
            pdf_data = {
                'pdf_file': pdf_file,
                'programme': programme,
                'coursework_type': extract_coursework_type_from_filename(pdf_file),
                'display_name': format_pdf_display_name(pdf_file)
            }
            
            student_database['student_modules'][student_id][module].append(pdf_data)
            student_database['students'][student_id]['modules'][module] = student_database['student_modules'][student_id][module]
            
            # Add to programme modules
            if programme not in student_database['programme_modules']:
                student_database['programme_modules'][programme] = {}
            
            if module not in student_database['programme_modules'][programme]:
                student_database['programme_modules'][programme][module] = []
            
            # Check if this PDF is already in the programme module list
            if not any(p['pdf_file'] == pdf_file for p in student_database['programme_modules'][programme][module]):
                student_database['programme_modules'][programme][module].append(pdf_data)
        
        logger.info(f"Loaded {len(student_database['students'])} students from database")
        
        # Log modules with multiple PDFs
        for student_id, modules in student_database['student_modules'].items():
            for module_name, pdfs in modules.items():
                if len(pdfs) > 1:
                    logger.info(f"Student {student_id}, Module '{module_name}' has {len(pdfs)} PDFs: {[p['pdf_file'] for p in pdfs]}")
        
        return student_database, "Database loaded successfully"
        
    except Exception as e:
        logger.error(f"Error loading student database: {e}")
        return None, f"Error loading database: {str(e)}"

def validate_student_credentials(student_id: str, code: str) -> Tuple[bool, Optional[Dict], str]:
    """Validate student ID and code"""
    if not st.session_state.student_database:
        return False, None, "Student database not loaded"
    
    student_id = student_id.strip().upper()
    
    try:
        code = int(code.strip())
    except ValueError:
        return False, None, "Code must be a number"
    
    # Check if student exists
    if student_id not in st.session_state.student_database['students']:
        return False, None, f"Student ID '{student_id}' not found in database"
    
    # Check if code matches
    stored_code = st.session_state.student_database['student_codes'][student_id]
    if code != stored_code:
        return False, None, f"Invalid code for student {student_id}"
    
    # Return student data
    student_data = st.session_state.student_database['students'][student_id]
    return True, student_data, "Authentication successful"

def generate_audio_response(text: str, voice: str = None) -> Optional[bytes]:
    """
    Generate audio response using OpenAI TTS
    
    Args:
        text: Text to convert to speech
        voice: Voice to use (defaults to Config.TTS_VOICE)
    
    Returns:
        Audio bytes or None if failed
    """
    if not client:
        logger.error("OpenAI client not initialized")
        return None
    
    if not text or not text.strip():
        logger.error("No text provided for audio generation")
        return None
    
    # Clean text for TTS (remove markdown and excessive formatting)
    clean_text = clean_text_for_tts(text)
    
    # Use provided voice or default
    selected_voice = voice or st.session_state.get('selected_voice', Config.TTS_VOICE)
    
    try:
        # Generate audio using OpenAI TTS
        response = client.audio.speech.create(
            model=Config.TTS_MODEL,
            voice=selected_voice,
            input=clean_text,
            response_format="mp3"
        )
        
        # Return audio bytes
        return response.content
        
    except Exception as e:
        logger.error(f"Error generating audio: {e}")
        return None

def clean_text_for_tts(text: str) -> str:
    """
    Clean text for text-to-speech by removing markdown and formatting
    
    Args:
        text: Raw text with potential markdown
        
    Returns:
        Cleaned text suitable for TTS
    """
    if not text:
        return ""
    
    # Remove markdown formatting
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # *italic*
    text = re.sub(r'_(.*?)_', r'\1', text)        # _italic_
    
    # Remove headers
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    
    # Remove code blocks
    text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # Remove special characters and emojis for better TTS
    text = re.sub(r'[🔑📄📚⚠️❌✅🤖🙋📊💾⏱️🔧🗑️🔄🔍🚨📁🎓📋🆔🔐]', '', text)    
    # Clean up multiple spaces and line breaks
    text = re.sub(r'\n+', '. ', text)
    text = re.sub(r'\s+', ' ', text)
    
    # Ensure proper sentence ending
    text = text.strip()
    if text and not text.endswith(('.', '!', '?')):
        text += '.'
    
    return text

def create_audio_player(audio_bytes: bytes, key: str = None) -> str:
    """
    Create an HTML audio player with the audio data
    
    Args:
        audio_bytes: Audio data in bytes
        key: Unique key for the audio player
        
    Returns:
        HTML string for the audio player
    """
    if not audio_bytes:
        return ""
    
    # Encode audio as base64
    audio_base64 = base64.b64encode(audio_bytes).decode()
    
    # Create unique key if not provided
    if not key:
        key = f"audio_{int(time.time() * 1000)}"
    
    # HTML audio player with custom styling
    audio_html = f"""
    <div class="audio-player-container" style="margin: 10px 0;">
        <div class="audio-controls" style="
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border-radius: 25px;
            padding: 10px 20px;
            display: flex;
            align-items: center;
            gap: 15px;
            box-shadow: 0 4px 15px rgba(102,126,234,0.3);
        ">
            <div style="color: white; font-weight: 500; display: flex; align-items: center; gap: 8px;">
                🔊 <span style="font-size: 14px;">{t('audio_response', default='Audio Response')}</span>
            </div>
            <audio controls style="
                height: 35px;
                border-radius: 17px;
                outline: none;
                flex: 1;
                min-width: 200px;
            ">
                <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mpeg">
                {t('audio_not_supported', default='Your browser does not support audio playback.')}
            </audio>
        </div>
    </div>
    """
    
    return audio_html

def create_audio_folder():
    """Create audio folder if it doesn't exist"""
    audio_path = Path(Config.AUDIO_FOLDER)
    audio_path.mkdir(exist_ok=True)
    return audio_path

def render_voice_selector():
    """Render voice selector in sidebar"""
    st.markdown(f"### 🎤 {t('voice_settings', default='Voice Settings')}")
    
    # Audio toggle
    current_audio_state = st.session_state.get('audio_enabled', True)
    audio_enabled = st.checkbox(
        f"🔊 {t('enable_audio', default='Enable Audio Responses')}", 
        value=current_audio_state,
        help=t('audio_help', default='Toggle audio responses for accessibility')
    )
    st.session_state.audio_enabled = audio_enabled
    
    if audio_enabled:
        # Voice selection
        voice_options = list(Config.SUPPORTED_VOICES.keys())
        voice_labels = [Config.SUPPORTED_VOICES[voice] for voice in voice_options]
        
        current_voice = st.session_state.get('selected_voice', Config.TTS_VOICE)
        
        # Find current voice index
        try:
            current_index = voice_options.index(current_voice)
        except ValueError:
            current_index = 0
        
        selected_voice = st.selectbox(
            f"🎭 {t('select_voice', default='Select Voice')}", 
            options=voice_options,
            format_func=lambda x: Config.SUPPORTED_VOICES[x],
            index=current_index,
            help=t('voice_help', default='Choose the voice for audio responses')
        )
        
        st.session_state.selected_voice = selected_voice
        
        # Test voice button
        if st.button(f"🎵 {t('test_voice', default='Test Voice')}", type="secondary"):
            test_text = t('test_audio_text', default='Hello! This is how I will sound when reading responses to you.')
            
            with st.spinner(t('generating_audio', default='Generating audio...')):
                audio_bytes = generate_audio_response(test_text, selected_voice)
                
            if audio_bytes:
                audio_html = create_audio_player(audio_bytes, key="voice_test")
                st.markdown(audio_html, unsafe_allow_html=True)
                st.success(t('audio_ready', default='Audio ready!'))
            else:
                st.error(t('audio_error', default='Failed to generate audio'))
    else:
        st.info(t('audio_disabled', default='Audio responses are disabled'))
def read_document(file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Read PDF or DOCX document"""
        try:
            if file_path.suffix.lower() == '.pdf':
                return read_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return read_docx(file_path)
            else:
                return None, {'error': f'Unsupported file type: {file_path.suffix}'}
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return None, {'error': str(e)}

def read_pdf(file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
    """Read PDF file and extract metadata"""
    try:
        reader = PdfReader(str(file_path))
        text = ""
        total_pages = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
            except Exception as e:
                logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                continue
        
        metadata = {
            'total_pages': total_pages,
            'file_size': file_path.stat().st_size,
            'file_type': 'PDF',
            'word_count': len(text.split()) if text else 0,
            'character_count': len(text),
        }
        
        return text, metadata
        
    except Exception as e:
        logger.error(f"Error reading PDF {file_path.name}: {e}")
        return None, {'error': str(e)}

def read_docx(file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
    """Read DOCX file"""
    try:
        doc = Document(str(file_path))
        text_parts = []
        paragraph_count = 0
        
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text_parts.append(para_text)
                paragraph_count += 1
        
        # Extract text from tables
        table_count = len(doc.tables)
        for table_idx, table in enumerate(doc.tables):
            text_parts.append(f"\n--- Table {table_idx + 1} ---")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        
        metadata = {
            'paragraphs': paragraph_count,
            'tables': table_count,
            'file_size': file_path.stat().st_size,
            'file_type': 'Word Document',
            'word_count': len(full_text.split()) if full_text else 0,
            'character_count': len(full_text),
        }
        
        return full_text, metadata
        
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path.name}: {e}")
        return None, {'error': str(e)}

def load_document_for_module(module_data: Dict) -> Tuple[Optional[str], Dict[str, Any], str]:
    """Enhanced document loading with support for multiple PDFs"""
    try:
        combined_content = ""
        combined_metadata = {}
        loaded_files = []
        
        if module_data['pdf_file'] == 'multiple':
            # Load all PDFs for this module
            total_pages = 0
            total_words = 0
            total_size = 0
            
            for pdf_data in module_data['all_pdfs']:
                pdf_filename = pdf_data['pdf_file']
                pdf_path = Path(Config.DATA_FOLDER) / pdf_filename
                
                if pdf_path.exists():
                    content, metadata = read_document(pdf_path)
                    if content:
                        # Add section header for each document
                        section_header = f"\n\n{'='*60}\n📄 {pdf_data['display_name']} ({pdf_data['coursework_type']})\nFile: {pdf_data['pdf_file']}\n{'='*60}\n"
                        combined_content += section_header + content
                        
                        # Aggregate metadata
                        total_pages += metadata.get('total_pages', 0)
                        total_words += metadata.get('word_count', 0)
                        total_size += metadata.get('file_size', 0)
                        loaded_files.append(pdf_filename)
                    else:
                        logger.warning(f"Failed to load content from {pdf_filename}")
                else:
                    logger.warning(f"File not found: {pdf_filename}")
            
            combined_metadata = {
                'total_pages': total_pages,
                'total_words': total_words,
                'total_size': total_size,
                'loaded_files': loaded_files,
                'file_count': len(loaded_files),
                'file_type': 'Multiple Documents'
            }
            
            if combined_content:
                return combined_content, combined_metadata, f"Loaded {len(loaded_files)} documents for {module_data['module']}"
            else:
                return None, combined_metadata, f"No content could be loaded for {module_data['module']}"
        
        else:
            # Load single PDF
            pdf_filename = module_data['pdf_file']
            pdf_path = Path(Config.DATA_FOLDER) / pdf_filename
            
            if not pdf_path.exists():
                return None, {}, f"Document not found: {pdf_filename}"
            
            content, metadata = read_document(pdf_path)
            
            if content:
                return content, metadata, f"Loaded {pdf_filename} successfully"
            else:
                return None, metadata, f"Failed to extract content from {pdf_filename}"
                
    except Exception as e:
        logger.error(f"Error loading document for module: {e}")
        return None, {}, f"Error: {str(e)}"

def generate_ai_response(question: str, document_content: str, module_info: Dict) -> str:
    """Enhanced AI response generation with support for multiple documents"""
    if not client:
        return f"🔑 **{t('api_key_missing')}**"
    
    if not document_content:
        return f"📄 **No document content available for {module_info['module']}**"
    
    try:
        # Create enhanced context-aware prompt
        document_info = ""
        if module_info['pdf_file'] == 'multiple':
            document_info = f"Multiple documents loaded for {module_info['module']}:\n"
            for pdf_data in module_info['all_pdfs']:
                document_info += f"- {pdf_data['display_name']} ({pdf_data['coursework_type']}) - File: {pdf_data['pdf_file']}\n"
        else:
            document_info = f"Document: {module_info['display_name']} ({module_info['coursework_type']}) - File: {module_info['pdf_file']}"
        
        system_prompt = f"""You are an expert academic assistant for University of Roehampton students. You are helping with the module: "{module_info['module']}" from the {module_info['programme']} programme.

{document_info}

DOCUMENT CONTENT:
{document_content[:Config.MAX_CONTENT_LENGTH]}

INSTRUCTIONS:
- Answer questions based ONLY on the provided document content
- If multiple documents are provided, clearly indicate which document contains specific information using the format **[Source: Document Name]**
- When referencing content, use the file names shown above for clarity
- Be helpful and educational, explaining concepts clearly
- If information isn't in the document(s), say so clearly
- Provide specific references to sections when possible
- Help with coursework understanding, but don't do the work for the student
- Encourage critical thinking and learning
- Be supportive and encouraging

CONTEXT:
- Module: {module_info['module']}
- Programme: {module_info['programme']}
- Materials: {module_info['coursework_type']}

Remember: You are helping a Roehampton University student understand their coursework materials. Always cite your sources when multiple documents are available."""

        response = client.chat.completions.create(
            model=Config.MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            max_tokens=Config.MAX_TOKENS,
            temperature=Config.TEMPERATURE,
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        logger.error(f"Error generating AI response: {e}")
        return f"❌ **Error generating response: {str(e)}**"

def reset_conversation():
    """Reset conversation to welcome state"""
    # Reset conversation flow states
    st.session_state.conversation_step = 'welcome'
    
    # Reset authentication data
    st.session_state.student_id = None
    st.session_state.student_code = None
    st.session_state.student_data = None
    st.session_state.selected_path = None
    
    # Reset module and coursework selection
    st.session_state.available_modules = []
    st.session_state.selected_module = None
    st.session_state.selected_coursework = None
    st.session_state.current_document = None
    
    # Reset chat data
    st.session_state.messages = []
    st.session_state.audio_responses = {}
    
    # Reset error handling
    st.session_state.error_message = None
    st.session_state.retry_count = 0

     # Reset ethics-specific states
    st.session_state.selected_ethics_category = None
    if 'ethics_document' in st.session_state:
        del st.session_state.ethics_document
    
    # Reset ethics-specific states if they exist
    if 'selected_ethics_category' in st.session_state:
        st.session_state.selected_ethics_category = None

def render_progress_indicator():
    """Render progress indicator for guided flow"""
    steps = {
        'welcome': 1,
        'path_selection': 1,
        'student_id': 2,
        'code': 3,
        'module': 4,
        'coursework': 5,
        'chat': 6
    }
    
    current_step = steps.get(st.session_state.conversation_step, 1)
    total_steps = 6
    
    progress = current_step / total_steps
    
    st.progress(progress)
    st.caption(f"Step {current_step} of {total_steps}")

   

def render_welcome_screen():
    """Render welcome screen with official Roehampton University branding"""
    
    # Load logo from assets folder
    logo_base64 = load_logo_from_assets()
    
    if logo_base64:
        # Welcome screen with Roehampton logo
        st.markdown(f"""
        <div class="roehampton-header">
            <div class="logo-title-container">
                <img src="data:image/png;base64,{logo_base64}" alt="University of Roehampton Logo" class="roehampton-logo">
                <div>
                    <h1>University Assistant</h1>
                </div>
            </div>
            <p>Your intelligent academic companion for coursework and ethics guidance</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Welcome screen without logo (fallback)
        st.markdown("""
        <div class="roehampton-header">
            <h1>🎓 University of Roehampton Assistant</h1>
            <p>Your intelligent academic companion for coursework and ethics guidance</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("### How can I help you today?")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📋 Ethics Document Help", 
                    help="Get assistance with ethics-related documents and guidelines",
                    use_container_width=True,
                    type="primary",
                    key="welcome_ethics_btn"):  # ← ADD THIS KEY
            st.session_state.selected_path = 'ethics'
            st.session_state.conversation_step = 'student_id'
            st.rerun()
    
    with col2:
        if st.button("📚 University Coursework Help", 
                    help="Get help with your specific coursework materials",
                    use_container_width=True,
                    type="primary",
                    key="welcome_coursework_btn"):  # ← ADD THIS KEY
            st.session_state.selected_path = 'coursework'
            st.session_state.conversation_step = 'student_id'
            st.rerun()
    
    # Feature highlights with Roehampton branding
    st.markdown("---")
    
    st.markdown("""
    <div class="feature-grid">
        <div class="feature-card">
            <div class="feature-icon">📋</div>
            <div class="feature-title">Ethics Guidance</div>
            <div class="feature-description">Access comprehensive ethics guidance based on university policies and the "Reforming Modernity" framework</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">📚</div>
            <div class="feature-title">Coursework Support</div>
            <div class="feature-description">Get personalized help with your module materials, assignments, and academic questions</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">🔐</div>
            <div class="feature-title">Secure Access</div>
            <div class="feature-description">Student authentication ensures you only access your own academic materials and information</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">🎤</div>
            <div class="feature-title">Audio Support</div>
            <div class="feature-description">Listen to responses with text-to-speech functionality for enhanced accessibility</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # col1, col2 = st.columns(2)
    
    # with col1:
    #     if st.button("📋 Ethics Document Help", 
    #                 help="Get assistance with ethics-related documents and guidelines",
    #                 use_container_width=True,
    #                 type="primary"):
    #         st.session_state.selected_path = 'ethics'
    #         st.session_state.conversation_step = 'student_id'
    #         st.rerun()
    
    # with col2:
    #     if st.button("📚 University Coursework Help", 
    #                 help="Get help with your specific coursework materials",
    #                 use_container_width=True,
    #                 type="primary"):
    #         st.session_state.selected_path = 'coursework'
    #         st.session_state.conversation_step = 'student_id'
    #         st.rerun()
    
    # Additional information
    st.markdown("---")
    st.markdown("""
    **What you can do:**
    - 📋 **Ethics Documents**: Access university ethics guidelines and policies
    - 📚 **Coursework Help**: Get assistance with your enrolled modules and assignments
    - 🔐 **Secure Access**: Authentication ensures you only see your own materials
    - 🎤 **Audio Support**: Listen to responses with text-to-speech functionality
    """)

def render_student_id_input():
    """Render student ID input screen"""
    st.markdown(f"### 🆔 Step 2: Enter Your Student ID")
    st.markdown(f"Please enter your Roehampton University Student ID to continue with **{st.session_state.selected_path}** assistance.")
    
    # Show error if exists
    if st.session_state.error_message:
        st.error(st.session_state.error_message)
        if st.session_state.retry_count > 2:
            st.warning("Having trouble? Please contact IT support or check your credentials.")
    
    student_id = st.text_input(
        "Student ID:",
        placeholder="e.g., A00034131",
        help="Enter your complete Roehampton University Student ID",
        key="student_id_input" 
    )
    
    col1, col2 = st.columns([1, 3])
    
    with col1:
        if st.button("🔙 Back", 
                    type="secondary",
                    key="student_id_back_btn"):  # ← ADD THIS KEY
            st.session_state.conversation_step = 'welcome'
            st.session_state.error_message = None
            st.session_state.retry_count = 0
            st.rerun()
    
    with col2:
        if st.button("Next ➡️", 
                    type="primary", 
                    disabled=not student_id.strip(),
                    key="student_id_next_btn"):  # ← ADD THIS KEY
            if student_id.strip():
                st.session_state.student_id = student_id.strip().upper()
                st.session_state.conversation_step = 'code'
                st.session_state.error_message = None
                st.rerun()

def render_code_input():
    """Render access code input screen"""
    st.markdown(f"### 🔐 Step 3: Enter Your Access Code")
    st.markdown(f"Student ID: **{st.session_state.student_id}**")
    st.markdown("Please enter your unique access code to verify your identity.")
    
    # Show error if exists
    if st.session_state.error_message:
        st.error(st.session_state.error_message)
    
    code = st.text_input(
        "Access Code:",
        type="password",
        placeholder="Enter your unique code",
        help="Enter the numerical code provided to you",
        key="access_code_input"
    )
    
    col1, col2 = st.columns([1, 3])
    
    with col1:
        if st.button("🔙 Back", 
                    type="secondary",
                    key="code_back_btn"):  # ← ADD THIS KEY
            st.session_state.conversation_step = 'student_id'
            st.session_state.error_message = None
            st.rerun()
    
    with col2:
        if st.button("Verify ✅", 
                    type="primary", 
                    disabled=not code.strip(),
                    key="code_verify_btn"):  # ← ADD THIS KEY
            if code.strip():
                # Validate credentials
                is_valid, student_data, message = validate_student_credentials(
                    st.session_state.student_id, 
                    code
                )
                
                if is_valid:
                    st.session_state.student_code = code
                    st.session_state.student_data = student_data
                    st.session_state.available_modules = student_data['modules']
                    st.session_state.error_message = None
                    st.session_state.retry_count = 0
                    
                    if st.session_state.selected_path == 'ethics':
                        st.session_state.conversation_step = 'ethics_chat'
                    else:
                        st.session_state.conversation_step = 'module'
                    
                    st.success(f"✅ Welcome, {st.session_state.student_id}!")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.session_state.error_message = message
                    st.session_state.retry_count += 1
                    st.rerun()
                

def get_module_selection_css():
    """Additional CSS for the redesigned module selection"""
    return """
    <style>
    /* Module Selection Specific Styles */
    .module-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(350px, 1fr));
        gap: 1.5rem;
        margin: 2rem 0;
    }
    
    .module-card {
        background: linear-gradient(145deg, #ffffff, #f8fffe);
        border: 2px solid #e8f5f0;
        border-radius: 16px;
        padding: 0;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(0, 168, 107, 0.1);
        overflow: hidden;
        position: relative;
    }
    
    .module-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 25px rgba(0, 168, 107, 0.2);
        border-color: #00a86b;
    }
    
    .module-card-header {
        background: linear-gradient(135deg, #00a86b, #008756);
        color: white;
        padding: 1.5rem;
        position: relative;
        overflow: hidden;
    }
    
    .module-card-header::before {
        content: '';
        position: absolute;
        top: -50%;
        right: -50%;
        width: 100px;
        height: 100px;
        background: radial-gradient(circle, rgba(255,255,255,0.2) 0%, transparent 70%);
        border-radius: 50%;
    }
    
    .module-title {
        font-size: 1.4rem;
        font-weight: 700;
        margin: 0 0 0.5rem 0;
        position: relative;
        z-index: 1;
    }
    
    .module-subtitle {
        font-size: 0.95rem;
        opacity: 0.9;
        margin: 0;
        position: relative;
        z-index: 1;
    }
    
    .module-card-body {
        padding: 1.5rem;
    }
    
    .pdf-list {
        margin: 0;
        padding: 0;
        list-style: none;
    }
    
    .pdf-item {
        background: #f8fffe;
        border: 1px solid #e8f5f0;
        border-radius: 10px;
        padding: 1rem;
        margin-bottom: 1rem;
        transition: all 0.2s ease;
        position: relative;
    }
    
    .pdf-item:hover {
        background: #f0fdf9;
        border-color: #00a86b;
        transform: translateX(5px);
    }
    
    .pdf-name {
        font-weight: 600;
        color: #00a86b;
        margin: 0 0 0.25rem 0;
        font-size: 1.05rem;
    }
    
    .pdf-meta {
        font-size: 0.85rem;
        color: #666;
        margin: 0;
        display: flex;
        align-items: center;
        gap: 1rem;
    }
    
    .pdf-badge {
        background: linear-gradient(135deg, #00a86b, #008756);
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 15px;
        font-size: 0.8rem;
        font-weight: 500;
    }
    
    .select-all-option {
        background: linear-gradient(135deg, #f0fdf9, #e8f5f0);
        border: 2px dashed #00a86b;
        border-radius: 12px;
        padding: 1.25rem;
        margin-top: 1rem;
        text-align: center;
        transition: all 0.3s ease;
    }
    
    .select-all-option:hover {
        background: linear-gradient(135deg, #e8f5f0, #d5f4e6);
        border-style: solid;
    }
    
    .select-all-title {
        font-weight: 600;
        color: #00a86b;
        margin: 0 0 0.5rem 0;
        font-size: 1.1rem;
    }
    
    .select-all-desc {
        color: #666;
        margin: 0;
        font-size: 0.9rem;
    }
    
    .action-buttons {
        display: flex;
        gap: 0.75rem;
        margin-top: 1rem;
        flex-wrap: wrap;
    }
    
    .btn-select {
        background: linear-gradient(135deg, #00a86b, #008756);
        color: white;
        border: none;
        padding: 0.6rem 1.2rem;
        border-radius: 8px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.2s ease;
        font-size: 0.9rem;
        flex: 1;
        min-width: 100px;
    }
    
    .btn-select:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 168, 107, 0.3);
    }
    
    .btn-select-all {
        background: linear-gradient(135deg, #1e3a5f, #2c3e50);
        color: white;
        border: none;
        padding: 0.6rem 1.2rem;
        border-radius: 8px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.2s ease;
        font-size: 0.9rem;
        width: 100%;
    }
    
    .btn-select-all:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(30, 58, 95, 0.3);
    }
    
    .single-module-card {
        background: linear-gradient(145deg, #ffffff, #f8fffe);
        border: 2px solid #e8f5f0;
        border-radius: 16px;
        padding: 0;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(0, 168, 107, 0.1);
        overflow: hidden;
        margin-bottom: 1.5rem;
    }
    
    .single-module-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(0, 168, 107, 0.2);
        border-color: #00a86b;
    }
    
    .progress-header {
        background: linear-gradient(135deg, #f8fffe, #f0fdf9);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        border: 1px solid #e8f5f0;
    }
    
    .progress-title {
        color: #00a86b;
        font-size: 1.8rem;
        font-weight: 700;
        margin: 0 0 0.5rem 0;
    }
    
    .progress-subtitle {
        color: #666;
        margin: 0 0 1rem 0;
        font-size: 1.05rem;
    }
    
    .student-info {
        display: flex;
        gap: 2rem;
        flex-wrap: wrap;
        align-items: center;
    }
    
    .info-badge {
        background: #00a86b;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.9rem;
    }
    
    /* Responsive Design */
    @media (max-width: 768px) {
        .module-grid {
            grid-template-columns: 1fr;
            gap: 1rem;
        }
        
        .student-info {
            flex-direction: column;
            gap: 1rem;
            align-items: flex-start;
        }
        
        .action-buttons {
            flex-direction: column;
        }
        
        .btn-select {
            flex: none;
            width: 100%;
        }
    }
    
    /* Animation for card entrance */
    @keyframes slideInUp {
        from {
            transform: translateY(30px);
            opacity: 0;
        }
        to {
            transform: translateY(0);
            opacity: 1;
        }
    }
    
    .module-card, .single-module-card {
        animation: slideInUp 0.5s ease forwards;
    }
    
    .module-card:nth-child(2) { animation-delay: 0.1s; }
    .module-card:nth-child(3) { animation-delay: 0.2s; }
    .module-card:nth-child(4) { animation-delay: 0.3s; }
    </style>
    """

def render_module_selection():
    """Redesigned module selection screen with modern card-based interface"""
    
    # Add module-specific CSS
    st.markdown(get_module_selection_css(), unsafe_allow_html=True)
    
    # Progress header
    st.markdown(f"""
    <div class="progress-header">
        <div class="progress-title">📚 Select Your Module</div>
        <div class="progress-subtitle">Choose the module you need assistance with</div>
        <div class="student-info">
            <div class="info-badge">👤 {st.session_state.student_id}</div>
            <div class="info-badge">🎓 {st.session_state.student_data['programme']}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Get modules from the database
    student_modules = st.session_state.student_database['student_modules'][st.session_state.student_id]
    
    if not student_modules:
        st.error("❌ No modules found for your account. Please contact support.")
        return
    
    # Create module cards
    modules_html = '<div class="module-grid">'
    
   
    
    modules_html += '</div>'
    
    # Display the HTML
    st.markdown(modules_html, unsafe_allow_html=True)
    
    # Now add the actual Streamlit buttons
    st.markdown("### Select Your Module:")
    
    for module_name, pdfs in student_modules.items():
        if len(pdfs) > 1:
            # Multi-PDF module buttons
            st.markdown(f"**📖 {module_name}** ({len(pdfs)} documents)")
            
            # Individual PDF buttons
            cols = st.columns(len(pdfs))
            for i, pdf_data in enumerate(pdfs):
                with cols[i]:
                    if st.button(
                        f"📄 {pdf_data['coursework_type']}", 
                        key=f"{module_name}_pdf_{i}_{pdf_data['pdf_file']}", 
                        help=f"Select {pdf_data['display_name']}",
                        use_container_width=True
                    ):
                        st.session_state.selected_module = {
                            'module': module_name,
                            'programme': pdf_data['programme'],
                            'pdf_file': pdf_data['pdf_file'],
                            'coursework_type': pdf_data['coursework_type'],
                            'display_name': pdf_data['display_name'],
                            'is_multi_pdf': True,
                            'all_pdfs': pdfs
                        }
                        st.session_state.conversation_step = 'coursework'
                        st.rerun()
            
            # Select all button
            if st.button(
                f"📚 All {module_name} Materials", 
                key=f"{module_name}_all_pdfs",
                help=f"Load all {len(pdfs)} documents together",
                type="secondary",
                use_container_width=True
            ):
                st.session_state.selected_module = {
                    'module': module_name,
                    'programme': pdfs[0]['programme'],
                    'pdf_file': 'multiple',
                    'coursework_type': 'All Materials',
                    'display_name': f"All {module_name} Materials",
                    'is_multi_pdf': True,
                    'all_pdfs': pdfs
                }
                st.session_state.conversation_step = 'coursework'
                st.rerun()
            
            st.markdown("---")
        
        else:
            # Single PDF module button
            pdf_data = pdfs[0]
            st.markdown(f"**📖 {module_name}**")
            
            if st.button(
                f"Select {module_name}", 
                key=f"{module_name}_single",
                help=f"Access {pdf_data['pdf_file']}",
                use_container_width=True,
                type="primary"
            ):
                st.session_state.selected_module = {
                    'module': module_name,
                    'programme': pdf_data['programme'],
                    'pdf_file': pdf_data['pdf_file'],
                    'coursework_type': pdf_data.get('coursework_type', 'Course Materials'),
                    'display_name': pdf_data.get('display_name', module_name),
                    'is_multi_pdf': False,
                    'all_pdfs': pdfs
                }
                st.session_state.conversation_step = 'coursework'
                st.rerun()
            
            st.markdown("---")
    
    # Back button
    st.markdown("### Navigation")
    if st.button("🔙 Back to Authentication", 
                type="secondary",
                key="module_back_btn",
                use_container_width=True):
        st.session_state.conversation_step = 'code'
        st.rerun()

def render_coursework_selection():
    """Render coursework selection screen"""
    st.markdown(f"### 📋 Step 5: Coursework Assistance")
    st.markdown(f"Module: **{st.session_state.selected_module['module']}**")
    st.markdown("What type of coursework help do you need?")
    
    # Coursework options
    coursework_options = [
        {
            'title': 'Assignment Questions',
            'description': 'Help understanding assignment requirements and questions',
            'type': 'assignment'
        },
        {
            'title': 'Reading Materials',
            'description': 'Assistance with course readings and materials',
            'type': 'reading'
        },
        {
            'title': 'Concepts & Theory',
            'description': 'Explanation of key concepts and theories',
            'type': 'concepts'
        },
        {
            'title': 'Exam Preparation',
            'description': 'Help preparing for examinations',
            'type': 'exam'
        },
        {
            'title': 'General Questions',
            'description': 'Any other questions about the module',
            'type': 'general'
        }
    ]
    
    for option in coursework_options:
        if st.button(
            f"📝 {option['title']}", 
            help=option['description'],
            use_container_width=True,
            key=f"coursework_{option['type']}"
        ):
            st.session_state.selected_coursework = option
            st.session_state.conversation_step = 'chat'
            st.rerun()
    
    # Back button
    if st.button("🔙 Back to Modules", type="secondary"):
        st.session_state.conversation_step = 'module'
        st.rerun()

def render_chat_interface():
    """Render the chat interface for the selected module/coursework"""
    
    if st.session_state.selected_path == 'ethics':
        # Handle ethics path
         if ETHICS_AVAILABLE:
            render_ethics_chat_interface()
         else:
            st.error("Ethics assistance is not available.")
            st.info("Please ensure 'reforming_modernity.pdf' is in your data folder and ethics_handler.py is properly configured.")
         return
    
    # Handle coursework path
    # Load document if not already loaded
    if not st.session_state.current_document and st.session_state.selected_module:
        with st.spinner("Loading your module materials..."):
            content, metadata, message = load_document_for_module(st.session_state.selected_module)
            if content:
                st.session_state.current_document = {
                    'content': content,
                    'metadata': metadata,
                    'module': st.session_state.selected_module
                }
                st.success(message)
            else:
                st.error(message)
                return
    
    # Header for coursework
    st.markdown(f"### 📚 {st.session_state.selected_module['module']}")
    st.markdown(f"**Coursework Type:** {st.session_state.selected_coursework['title']}")
    st.markdown(f"**Programme:** {st.session_state.student_data['programme']}")
    
    # Coursework-specific examples
    with st.expander("💡 Example Questions", expanded=False):
        coursework_type = st.session_state.selected_coursework['type']
        if coursework_type == 'assignment':
            st.markdown("""
            - "What are the key requirements for this assignment?"
            - "How should I structure my report?"
            - "What citation format should I use?"
            - "What are the assessment criteria?"
            """)
        elif coursework_type == 'reading':
            st.markdown("""
            - "Can you summarize the main concepts in this module?"
            - "What are the key theories I should understand?"
            - "Which readings are most important for the exam?"
            - "How do these concepts relate to practical applications?"
            """)
        elif coursework_type == 'concepts':
            st.markdown("""
            - "Can you explain [specific concept] in simple terms?"
            - "How does [theory A] relate to [theory B]?"
            - "What are some real-world examples of this concept?"
            - "Why is this concept important in the field?"
            """)
        elif coursework_type == 'exam':
            st.markdown("""
            - "What topics are likely to be on the exam?"
            - "How should I prepare for this type of assessment?"
            - "Can you create practice questions for me?"
            - "What are the key points I should remember?"
            """)
        else:
            st.markdown("""
            - "What are the learning objectives for this module?"
            - "How can I improve my understanding of this subject?"
            - "What additional resources do you recommend?"
            - "How does this module connect to my overall programme?"
            """)
    
    # Chat messages with audio support
    for i, message in enumerate(st.session_state.messages):
        message_key = f"msg_{i}_{message.get('timestamp', time.time())}"
        
        if message["role"] == "user":
            st.markdown(f"""
            <div style="background: #e3f2fd; color: #000; padding: 1rem; border-radius: 10px; margin: 1rem 0; border-left: 4px solid #2196f3;">
                <strong>🙋 You:</strong><br>{message["content"]}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div style="background: #f1f8e9; color: #000; padding: 1rem; border-radius: 10px; margin: 1rem 0; border-left: 4px solid #4caf50;">
                <strong>🤖 Course Assistant:</strong><br>{message["content"]}
            </div>
            """, unsafe_allow_html=True)
            
            # Add audio player if audio is enabled
            if st.session_state.get('audio_enabled', True):
                # Check if we already have audio for this message
                if message_key not in st.session_state.audio_responses:
                    # Generate audio for this message
                    with st.spinner(t('generating_audio', default='Generating audio...')):
                        audio_bytes = generate_audio_response(
                            message["content"], 
                            st.session_state.get('selected_voice', Config.TTS_VOICE)
                        )
                        if audio_bytes:
                            st.session_state.audio_responses[message_key] = audio_bytes
                
                # Display audio player if we have audio
                if message_key in st.session_state.audio_responses:
                    audio_html = create_audio_player(
                        st.session_state.audio_responses[message_key], 
                        key=message_key
                    )
                    st.markdown(audio_html, unsafe_allow_html=True)
    
    # Chat input
    if prompt := st.chat_input("Ask me about your coursework..."):
        # Add user message
        st.session_state.messages.append({
            "role": "user",
            "content": prompt,
            "timestamp": time.time()
        })
        
        # Generate AI response for coursework
        with st.spinner("Analyzing your coursework materials..."):
            response = generate_ai_response(
                prompt,
                st.session_state.current_document['content'],
                st.session_state.current_document['module']
            )
        
        # Add AI response
        ai_message = {
            "role": "assistant",
            "content": response,
            "timestamp": time.time()
        }
        st.session_state.messages.append(ai_message)
        
        # Pre-generate audio if enabled (for better UX)
        if st.session_state.get('audio_enabled', True) and response:
            message_key = f"msg_{len(st.session_state.messages)-1}_{ai_message['timestamp']}"
            try:
                with st.spinner(t('preparing_audio', default='Preparing audio...')):
                    audio_bytes = generate_audio_response(
                        response, 
                        st.session_state.get('selected_voice', Config.TTS_VOICE)
                    )
                    if audio_bytes:
                        st.session_state.audio_responses[message_key] = audio_bytes
            except Exception as e:
                logger.error(f"Error pre-generating audio: {e}")
        
        st.rerun()
    
    # Control buttons
    col1, col2, col3 = st.columns([1, 1, 2])
    
    with col1:
        if st.button("🔄 New Session", type="secondary"):
            reset_conversation()
            st.rerun()
    
    with col2:
        if st.button("🔙 Change Module", type="secondary"):
            st.session_state.conversation_step = 'module'
            st.session_state.messages = []
            st.session_state.current_document = None
            st.rerun()

def render_sidebar():
    """Render sidebar with student info and controls"""
    with st.sidebar:
        # Language selector
        st.markdown(f"### 🌐 {t('language_selector')}")
        render_language_selector()
        # Voice settings
        render_voice_selector()
        
        st.markdown("---")
        
        # Student information (if authenticated)
        if st.session_state.student_id and st.session_state.student_data:
            st.markdown("### 👤 Student Information")
            st.markdown(f"""
            <div style="background: #f0f2f6; color: #000; padding: 1rem; border-radius: 8px;">
                <p><strong>ID:</strong> {st.session_state.student_id}</p>
                <p><strong>Programme:</strong> {st.session_state.student_data['programme']}</p>
                <p><strong>Modules:</strong> {len(st.session_state.available_modules)}</p>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("---")
        
        # Current session info
        if st.session_state.conversation_step != 'welcome':
            st.markdown("### 📍 Current Session")
            st.markdown(f"**Path:** {st.session_state.selected_path or 'Not selected'}")
            
            # ADD ETHICS-SPECIFIC INFO
            if st.session_state.selected_path == 'ethics':
                st.markdown("**Document:** Reforming Modernity")
                if st.session_state.selected_ethics_category:
                    st.markdown(f"**Category:** {st.session_state.selected_ethics_category['title']}")
            elif st.session_state.selected_module:
                st.markdown(f"**Module:** {st.session_state.selected_module['module']}")
                if st.session_state.selected_coursework:
                    st.markdown(f"**Type:** {st.session_state.selected_coursework['title']}")
            
            st.markdown("---")
        
        # Database status
        st.markdown("### 📊 System Status")
        if st.session_state.database_loaded:
            st.success("✅ Database Connected")
            total_students = len(st.session_state.student_database['students'])
            total_programmes = len(st.session_state.student_database['programme_modules'])
            st.markdown(f"**Students:** {total_students}")
            st.markdown(f"**Programmes:** {total_programmes}")
        else:
            st.error("❌ Database Not Loaded")
        
        if OPENAI_API_KEY:
            st.success("✅ AI Service Connected")
        else:
            st.error("❌ AI Service Not Available")
        
        st.markdown("---")
        
        # Quick actions
        st.markdown("### ⚡ Quick Actions")
        
        if st.button("🏠 Start Over", use_container_width=True, type="secondary"):
            reset_conversation()
            st.rerun()
        
        if st.session_state.conversation_step == 'chat':
            if st.button("🗑️ Clear Chat", use_container_width=True, type="secondary"):
                st.session_state.messages = []
                st.session_state.audio_responses = {}
                st.rerun()

def get_enhanced_css() -> str:    
    """Get enhanced CSS with official Roehampton University brand colors"""
    base_css = """
    <style>
        /* Import Google Fonts */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Arabic:wght@300;400;500;600;700&display=swap');
        
        /* Roehampton University Brand Colors */
        :root {
            --roehampton-green: #00A86B;        /* Primary Roehampton Green */
            --roehampton-dark-green: #008756;   /* Darker shade for hover states */
            --roehampton-light-green: #33BA85;  /* Lighter shade for accents */
            --roehampton-navy: #1E3A5F;         /* Navy from logo text */
            --roehampton-charcoal: #2C3E50;     /* Dark text color */
            --success-color: #00A86B;           /* Use Roehampton green for success */
            --primary-color: #00A86B;           /* Primary brand color */
            --secondary-color: #1E3A5F;         /* Secondary navy color */
            --error-color: #E74C3C;
            --warning-color: #F39C12;
            --info-color: #3498DB;
            --background-light: #F8FFFE;        /* Very light green tint */
            --background-white: #FFFFFF;
            --text-primary: #2C3E50;
            --text-secondary: #7F8C8D;
            --border-color: #E8F5F0;            /* Light green border */
            --shadow: 0 2px 4px rgba(0, 168, 107, 0.1);
            --shadow-lg: 0 8px 25px rgba(0, 168, 107, 0.15);
        }
        
        /* Global font with multi-language support */
        .main, .sidebar .sidebar-content {
            font-family: 'Inter', 'Noto Sans Arabic', 'Arial', sans-serif !important;
            background-color: var(--background-light);
        }
        
        /* Arabic text specific styling */
        [lang="ar"], .arabic-text {
            font-family: 'Noto Sans Arabic', 'Arial', 'Tahoma', sans-serif !important;
            line-height: 1.8 !important;
            text-align: right !important;
        }
        
        /* Roehampton University Branded Header */
        .roehampton-header {
            background: linear-gradient(135deg, var(--roehampton-green), var(--roehampton-dark-green));
            padding: 2.5rem 2rem;
            border-radius: 15px;
            margin-bottom: 2rem;
            color: white;
            text-align: center;
            box-shadow: var(--shadow-lg);
            position: relative;
            overflow: hidden;
        }
        
        .roehampton-header::before {
            content: '';
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
            animation: shimmer 3s infinite;
        }
        
        @keyframes shimmer {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        
        .logo-title-container {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 2rem;
            margin-bottom: 1rem;
            position: relative;
            z-index: 1;
        }
        
        .roehampton-logo {
            height: 90px;
            width: auto;
            background: white;
            padding: 0.75rem;
            border-radius: 12px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        }
        
        .roehampton-header h1 {
            margin: 0;
            font-weight: 700;
            font-size: 2.8rem;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }
        
        .roehampton-header p {
            margin: 0;
            opacity: 0.95;
            font-size: 1.3rem;
            font-weight: 400;
        }
        
        /* Branded Buttons */
        .stButton > button {
            border-radius: 12px;
            font-weight: 600;
            transition: all 0.3s ease;
            border: none;
            font-size: 1.1rem;
            padding: 0.8rem 1.5rem;
        }
        
        .stButton > button[data-baseweb="button"][kind="primary"] {
            background: linear-gradient(135deg, var(--roehampton-green), var(--roehampton-dark-green));
            color: white;
            box-shadow: var(--shadow);
        }
        
        .stButton > button[data-baseweb="button"][kind="primary"]:hover {
            transform: translateY(-2px);
            box-shadow: var(--shadow-lg);
            background: linear-gradient(135deg, var(--roehampton-dark-green), var(--roehampton-green));
        }
        
        .stButton > button[data-baseweb="button"][kind="secondary"] {
            background: white;
            color: var(--roehampton-green);
            border: 2px solid var(--roehampton-green);
        }
        
        .stButton > button[data-baseweb="button"][kind="secondary"]:hover {
            background: var(--roehampton-green);
            color: white;
            transform: translateY(-1px);
        }
        
        /* Chat message containers with Roehampton branding */
        .chat-message {
            padding: 1.5rem;
            border-radius: 15px;
            margin-bottom: 1.5rem;
            box-shadow: var(--shadow);
            transition: transform 0.2s ease;
            border: 1px solid var(--border-color);
        }
        
        .chat-message:hover {
            transform: translateY(-2px);
            box-shadow: var(--shadow-lg);
        }
        
        .user-message {
            background: linear-gradient(135deg, #E8F8F5, #D5F4E6);
            border-left: 4px solid var(--roehampton-green);
            margin-left: 2rem;
        }
        
        .assistant-message {
            background: linear-gradient(135deg, #F8FFFE, #F0FDF9);
            border-left: 4px solid var(--roehampton-light-green);
            margin-right: 2rem;
        }
        
        .ethics-message {
            background: linear-gradient(135deg, #F0F4F8, #E2E8F0);
            border-left: 4px solid var(--roehampton-navy);
            margin-right: 2rem;
        }
        
        .message-header {
            font-weight: 600;
            margin-bottom: 0.75rem;
            color: var(--text-primary);
            display: flex;
            align-items: center;
            gap: 0.5rem;
            font-size: 1rem;
        }
        
        .message-content {
            color: var(--text-primary);
            line-height: 1.7;
            font-size: 1rem;
        }
        
        /* Progress indicator with Roehampton colors */
        .stProgress > div > div > div {
            background: linear-gradient(90deg, var(--roehampton-green), var(--roehampton-light-green));
        }
        
        /* Input styling with Roehampton theme */
        .stTextInput > div > div > input {
            border-radius: 10px;
            border: 2px solid var(--border-color);
            padding: 0.75rem 1rem;
            font-size: 1rem;
            transition: all 0.2s ease;
        }
        
        .stTextInput > div > div > input:focus {
            border-color: var(--roehampton-green);
            box-shadow: 0 0 0 3px rgba(0, 168, 107, 0.1);
        }
        
        /* Audio player with Roehampton branding */
        .audio-player-container {
            margin: 15px 0;
            padding: 0;
        }
        
        .audio-controls {
            background: linear-gradient(135deg, var(--roehampton-green), var(--roehampton-dark-green));
            border-radius: 25px;
            padding: 12px 20px;
            display: flex;
            align-items: center;
            gap: 15px;
            box-shadow: var(--shadow-lg);
            transition: all 0.3s ease;
        }
        
        .audio-controls:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 25px rgba(0, 168, 107, 0.3);
        }
        
        .audio-controls audio {
            height: 35px;
            border-radius: 17px;
            outline: none;
            flex: 1;
            min-width: 200px;
            background: rgba(255,255,255,0.2);
        }
        
        /* Sidebar styling with Roehampton theme */
        .sidebar-section {
            background: var(--background-white);
            padding: 1.5rem;
            border-radius: 12px;
            margin-bottom: 1rem;
            border: 1px solid var(--border-color);
            box-shadow: var(--shadow);
        }
        
        .sidebar-section h3 {
            color: var(--roehampton-green);
            margin-top: 0;
        }
        
        .info-item {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 0.75rem 0;
            border-bottom: 1px solid var(--border-color);
        }
        
        .info-item:last-child {
            border-bottom: none;
        }
        
        .info-label {
            font-weight: 500;
            color: var(--text-secondary);
            font-size: 0.95rem;
        }
        
        .info-value {
            font-weight: 600;
            color: var(--roehampton-green);
            font-size: 0.95rem;
        }
        
        /* Status indicators with Roehampton branding */
        .status-success {
            background: linear-gradient(135deg, var(--roehampton-green), var(--roehampton-dark-green));
            color: white;
            padding: 0.6rem 1.2rem;
            border-radius: 25px;
            font-size: 0.9rem;
            font-weight: 600;
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
            box-shadow: var(--shadow);
        }
        
        .status-error {
            background: linear-gradient(135deg, var(--error-color), #C0392B);
            color: white;
            padding: 0.6rem 1.2rem;
            border-radius: 25px;
            font-size: 0.9rem;
            font-weight: 600;
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
        }
        
        .status-info {
            background: linear-gradient(135deg, var(--roehampton-navy), #34495E);
            color: white;
            padding: 0.6rem 1.2rem;
            border-radius: 25px;
            font-size: 0.9rem;
            font-weight: 600;
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
        }
        
        /* Module cards with Roehampton styling */
        .module-card {
            border: 2px solid var(--border-color);
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 1rem;
            background: var(--background-white);
            transition: all 0.3s ease;
        }
        
        .module-card:hover {
            border-color: var(--roehampton-green);
            box-shadow: var(--shadow-lg);
            transform: translateY(-2px);
        }
        
        /* Welcome page features */
        .feature-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 1.5rem;
            margin: 2rem 0;
        }
        
        .feature-card {
            background: var(--background-white);
            padding: 1.5rem;
            border-radius: 12px;
            border: 1px solid var(--border-color);
            box-shadow: var(--shadow);
            transition: all 0.3s ease;
        }
        
        .feature-card:hover {
            transform: translateY(-3px);
            box-shadow: var(--shadow-lg);
            border-color: var(--roehampton-light-green);
        }
        
        .feature-icon {
            font-size: 2.5rem;
            margin-bottom: 1rem;
            text-align: center;
        }
        
        .feature-title {
            font-size: 1.2rem;
            font-weight: 600;
            color: var(--roehampton-green);
            margin-bottom: 0.5rem;
            text-align: center;
        }
        
        .feature-description {
            color: var(--text-secondary);
            text-align: center;
            line-height: 1.5;
        }
        
        /* Responsive design */
        @media (max-width: 768px) {
            .roehampton-header h1 {
                font-size: 2rem;
            }
            
            .logo-title-container {
                flex-direction: column;
                gap: 1rem;
            }
            
            .roehampton-logo {
                height: 70px;
            }
            
            .chat-message {
                margin-left: 0.5rem;
                margin-right: 0.5rem;
                padding: 1rem;
            }
            
            .user-message, .assistant-message, .ethics-message {
                margin-left: 0;
                margin-right: 0;
            }
        }
        
        /* Loading animation with Roehampton colors */
        .loading-spinner {
            display: inline-block;
            width: 20px;
            height: 20px;
            border: 2px solid var(--border-color);
            border-radius: 50%;
            border-top-color: var(--roehampton-green);
            animation: spin 1s ease-in-out infinite;
        }
        
        @keyframes spin {
            to { transform: rotate(360deg); }
        }
        
        /* Streamlit specific overrides */
        .stAlert {
            border-radius: 12px;
            border: none;
            font-weight: 500;
        }
        
        .stExpander {
            border-radius: 10px;
            border: 1px solid var(--border-color);
        }
        
        .stSelectbox > div > div {
            border-radius: 8px;
            border-color: var(--border-color);
        }
        
        /* Custom checkbox styling */
        .stCheckbox > label {
            color: var(--text-primary);
        }
        
        .stCheckbox > label > div {
            background-color: var(--roehampton-green);
        }
    </style>
    """
    
    # Add RTL-specific CSS if needed
    from localization import get_rtl_css
    rtl_css = get_rtl_css()
    
    return base_css + rtl_css

def main():
    """Main application function with guided conversation flow"""
    try:
        # Initialize session state
        initialize_session_state()
        
        # Create audio folder
        create_audio_folder()
        
        # Apply enhanced CSS
        st.markdown(get_enhanced_css(), unsafe_allow_html=True)
        
        # Load student database if not loaded
        if not st.session_state.database_loaded:
            with st.spinner("Loading student database..."):
                database, message = load_student_database()
                if database:
                    st.session_state.student_database = database
                    st.session_state.database_loaded = True
                    logger.info("Student database loaded successfully")
                else:
                    st.error(f"Failed to load student database: {message}")
                    st.stop()
        
        # Render sidebar
        render_sidebar()
        
        # Render progress indicator (except for welcome screen)
        if st.session_state.conversation_step != 'welcome':
            render_progress_indicator()
            st.markdown("---")
        
        # Render appropriate screen based on conversation step
        if st.session_state.conversation_step == 'welcome':
            render_welcome_screen()
        
        elif st.session_state.conversation_step == 'student_id':
            render_student_id_input()
        
        elif st.session_state.conversation_step == 'code':
            render_code_input()
        
        elif st.session_state.conversation_step == 'module':
            render_module_selection()
        
        elif st.session_state.conversation_step == 'coursework':
            render_coursework_selection()
        
        elif st.session_state.conversation_step == 'ethics_selection':
            if ETHICS_AVAILABLE:
                render_ethics_chat_interface()
            else:
                st.error("Ethics assistance is not available. Please ensure ethics_handler.py is properly configured.")
                if st.button("🔙 Back"):
                    st.session_state.conversation_step = 'welcome'
                    st.rerun()
        
        elif st.session_state.conversation_step == 'ethics_chat':
            if ETHICS_AVAILABLE:
                render_ethics_chat_interface()
            else:
                st.error("Ethics assistance is not available. Please ensure ethics_handler.py is properly configured.")
                if st.button("🔙 Back"):
                    st.session_state.conversation_step = 'welcome'
                    st.rerun()
        
        elif st.session_state.conversation_step == 'chat':
            render_chat_interface()
        
        else:
            st.error("Unknown conversation step. Please restart.")
            if st.button("🔄 Restart"):
                reset_conversation()
                st.rerun()
    
    except Exception as e:
        logger.error(f"Application error: {e}")
        st.error(f"🚨 **Application Error**: {str(e)}")
        
        # Show detailed error information for debugging
        if st.checkbox("Show detailed error information (for debugging)"):
            import traceback
            st.code(traceback.format_exc())
        
        st.info("Please try the following:")
        st.markdown("""
        1. **Refresh the page** and try again
        2. **Check your .env file** - ensure OPENAI_API_KEY is set
        3. **Verify your Excel file** - ensure student_modules_with_pdfs.xlsx exists
        4. **Check file permissions** - ensure the app can read your files
        5. **Restart the application** if the problem persists
        """)
        
        if st.button("🔄 Reset Application"):
            # Clear all session state
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

if __name__ == '__main__':
    main()